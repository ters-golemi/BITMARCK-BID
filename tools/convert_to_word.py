#!/usr/bin/env python3
"""
Document Converter Script
Converts Markdown and Text files to Word documents (.docx)
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import markdown
import re
from pathlib import Path

class DocumentConverter:
    """Convert various document formats to Word documents"""
    
    def __init__(self, source_folder, output_folder=None):
        self.source_folder = Path(source_folder)
        self.output_folder = Path(output_folder) if output_folder else self.source_folder / "word_documents"
        
        # Create output directory if it doesn't exist
        self.output_folder.mkdir(exist_ok=True)
        
    def convert_all_documents(self):
        """Convert all supported documents in the folder"""
        
        converted_files = []
        
        # Get all files in the source folder
        for file_path in self.source_folder.iterdir():
            if file_path.is_file():
                file_extension = file_path.suffix.lower()
                
                if file_extension == '.md':
                    output_file = self.convert_markdown_to_word(file_path)
                    converted_files.append(output_file)
                elif file_extension == '.txt':
                    output_file = self.convert_text_to_word(file_path)
                    converted_files.append(output_file)
                else:
                    print(f"Skipping unsupported file: {file_path.name}")
        
        return converted_files
    
    def convert_markdown_to_word(self, markdown_file):
        """Convert Markdown file to Word document"""
        
        print(f"Converting Markdown file: {markdown_file.name}")
        
        # Read the markdown content with fallback encoding
        try:
            with open(markdown_file, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
        except UnicodeDecodeError:
            try:
                with open(markdown_file, 'r', encoding='latin-1') as f:
                    markdown_content = f.read()
            except UnicodeDecodeError:
                with open(markdown_file, 'r', encoding='cp1252') as f:
                    markdown_content = f.read()
        
        # Convert markdown to HTML first
        html_content = markdown.markdown(markdown_content, extensions=['extra', 'codehilite'])
        
        # Create Word document
        doc = Document()
        
        # Set document styles
        self.setup_document_styles(doc)
        
        # Parse and add content to Word document
        self.parse_markdown_to_word(markdown_content, doc)
        
        # Generate output filename
        output_filename = self.output_folder / f"{markdown_file.stem}.docx"
        
        # Save the document
        doc.save(output_filename)
        print(f"Saved: {output_filename}")
        
        return output_filename
    
    def convert_text_to_word(self, text_file):
        """Convert plain text file to Word document"""
        
        print(f"Converting text file: {text_file.name}")
        
        # Read the text content with fallback encoding
        try:
            with open(text_file, 'r', encoding='utf-8') as f:
                text_content = f.read()
        except UnicodeDecodeError:
            try:
                with open(text_file, 'r', encoding='latin-1') as f:
                    text_content = f.read()
            except UnicodeDecodeError:
                with open(text_file, 'r', encoding='cp1252') as f:
                    text_content = f.read()
        
        # Create Word document
        doc = Document()
        
        # Set document styles
        self.setup_document_styles(doc)
        
        # Add title based on filename
        title = text_file.stem.replace('_', ' ').replace('-', ' ').title()
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Parse and add content
        self.parse_text_to_word(text_content, doc)
        
        # Generate output filename
        output_filename = self.output_folder / f"{text_file.stem}.docx"
        
        # Save the document
        doc.save(output_filename)
        print(f"Saved: {output_filename}")
        
        return output_filename
    
    def setup_document_styles(self, doc):
        """Setup document styles and formatting"""
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def parse_markdown_to_word(self, markdown_content, doc):
        """Parse markdown content and add to Word document"""
        
        lines = markdown_content.split('\n')
        current_list = None
        code_block = False
        code_language = None
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                # Empty line - add paragraph break
                doc.add_paragraph()
                i += 1
                continue
            
            # Handle code blocks
            if line.startswith('```'):
                if not code_block:
                    # Start of code block
                    code_block = True
                    code_language = line[3:].strip() if len(line) > 3 else 'text'
                    i += 1
                    continue
                else:
                    # End of code block
                    code_block = False
                    i += 1
                    continue
            
            if code_block:
                # Inside code block - add as formatted code
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(line)
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(9)
                i += 1
                continue
            
            # Handle headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('#').strip()
                
                if level <= 3:
                    doc.add_heading(header_text, level=level)
                else:
                    # For levels > 3, use bold paragraph
                    para = doc.add_paragraph()
                    run = para.add_run(header_text)
                    run.bold = True
                    run.font.size = Pt(12)
                
                i += 1
                continue
            
            # Handle lists
            if line.startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\.\s', line):
                list_text = re.sub(r'^[-*+]\s*|^\d+\.\s*', '', line)
                para = doc.add_paragraph(list_text, style='List Bullet' if line.startswith(('- ', '* ', '+ ')) else 'List Number')
                i += 1
                continue
            
            # Handle tables (basic support)
            if '|' in line and not line.startswith('|'):
                # This might be a table row
                table_rows = []
                j = i
                while j < len(lines) and '|' in lines[j]:
                    table_rows.append(lines[j])
                    j += 1
                
                if len(table_rows) > 1:
                    self.add_table_to_word(table_rows, doc)
                    i = j
                    continue
            
            # Handle bold and italic text
            para = doc.add_paragraph()
            self.add_formatted_text_to_paragraph(line, para)
            
            i += 1
    
    def parse_text_to_word(self, text_content, doc):
        """Parse plain text content and add to Word document"""
        
        lines = text_content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                doc.add_paragraph()
                continue
            
            # Try to identify structure in plain text
            if line.isupper() and len(line) < 100:
                # Likely a section header
                doc.add_heading(line.title(), level=2)
            elif line.endswith(':') and len(line) < 80:
                # Likely a subsection
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.bold = True
            elif line.startswith(('- ', '* ', '• ')):
                # List item
                list_text = line.lstrip('- *• ')
                doc.add_paragraph(list_text, style='List Bullet')
            elif re.match(r'^\d+\.?\s', line):
                # Numbered list
                list_text = re.sub(r'^\d+\.?\s*', '', line)
                doc.add_paragraph(list_text, style='List Number')
            else:
                # Regular paragraph
                doc.add_paragraph(line)
    
    def add_formatted_text_to_paragraph(self, text, paragraph):
        """Add text with markdown formatting to paragraph"""
        
        # Simple regex patterns for markdown formatting
        patterns = [
            (r'\*\*(.*?)\*\*', 'bold'),
            (r'\*(.*?)\*', 'italic'),
            (r'`(.*?)`', 'code')
        ]
        
        current_text = text
        last_end = 0
        
        # Find all formatting matches
        all_matches = []
        for pattern, format_type in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                all_matches.append((match.start(), match.end(), match.group(1), format_type))
        
        # Sort matches by position
        all_matches.sort(key=lambda x: x[0])
        
        # Add text with formatting
        for start, end, match_text, format_type in all_matches:
            # Add text before the match
            if start > last_end:
                paragraph.add_run(text[last_end:start])
            
            # Add formatted text
            run = paragraph.add_run(match_text)
            if format_type == 'bold':
                run.bold = True
            elif format_type == 'italic':
                run.italic = True
            elif format_type == 'code':
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            
            last_end = end
        
        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])
    
    def add_table_to_word(self, table_rows, doc):
        """Add a table to the Word document"""
        
        # Parse table rows
        parsed_rows = []
        for row in table_rows:
            cells = [cell.strip() for cell in row.split('|') if cell.strip()]
            if cells:
                parsed_rows.append(cells)
        
        if not parsed_rows:
            return
        
        # Create table
        max_cols = max(len(row) for row in parsed_rows)
        table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
        table.style = 'Light Grid Accent 1'
        
        # Populate table
        for i, row_data in enumerate(parsed_rows):
            row_cells = table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                if j < len(row_cells):
                    row_cells[j].text = cell_data
                    
                    # Make header row bold
                    if i == 0:
                        for paragraph in row_cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

def main():
    """Main conversion function"""
    
    # Source folder containing documents to convert
    source_folder = "/Users/adiscato/Python/TENDER"
    output_folder = "/Users/adiscato/Python/TENDER/word_documents"
    
    print("=== TENDER Documents to Word Converter ===")
    print(f"Source folder: {source_folder}")
    print(f"Output folder: {output_folder}")
    print()
    
    # Initialize converter
    converter = DocumentConverter(source_folder, output_folder)
    
    # Convert all documents
    try:
        converted_files = converter.convert_all_documents()
        
        print("\n=== Conversion Complete ===")
        print(f"Successfully converted {len(converted_files)} files:")
        
        for file_path in converted_files:
            print(f"  - {file_path.name}")
        
        print(f"\nAll Word documents saved to: {output_folder}")
        
    except Exception as e:
        print(f"Error during conversion: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
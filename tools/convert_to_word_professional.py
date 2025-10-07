#!/usr/bin/env python3
"""
Enhanced Document Converter Script
Converts documents from TENDER folder to professional Word documents (.docx)
with improved formatting and styling
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import markdown
import re
from pathlib import Path
from datetime import datetime

class EnhancedDocumentConverter:
    """Enhanced converter for professional Word documents"""
    
    def __init__(self, source_folder, output_folder=None):
        self.source_folder = Path(source_folder)
        self.output_folder = Path(output_folder) if output_folder else self.source_folder / "professional_word_documents"
        
        # Create output directory if it doesn't exist
        self.output_folder.mkdir(exist_ok=True)
        
    def convert_all_documents_professional(self):
        """Convert all documents with professional formatting"""
        
        converted_files = []
        
        print("=== Enhanced Professional Document Conversion ===")
        print(f"Converting documents from: {self.source_folder}")
        print(f"Output directory: {self.output_folder}")
        print()
        
        # Get all files in the source folder
        for file_path in self.source_folder.iterdir():
            if file_path.is_file() and not file_path.name.startswith('.'):
                file_extension = file_path.suffix.lower()
                
                if file_extension == '.md':
                    output_file = self.convert_markdown_professional(file_path)
                    converted_files.append(output_file)
                elif file_extension == '.txt':
                    output_file = self.convert_text_professional(file_path)
                    converted_files.append(output_file)
                else:
                    print(f"Skipping unsupported file: {file_path.name}")
        
        return converted_files
    
    def convert_markdown_professional(self, markdown_file):
        """Convert Markdown file to professional Word document"""
        
        print(f"Converting Markdown: {markdown_file.name}")
        
        # Read content with encoding fallback
        content = self.read_file_with_encoding(markdown_file)
        
        # Create professional Word document
        doc = Document()
        self.setup_professional_styles(doc)
        
        # Add header/footer
        self.add_document_header_footer(doc, markdown_file.stem)
        
        # Add title page
        self.add_title_page(doc, markdown_file.stem, "Technical Documentation", "BITMARCK Tender Response")
        
        # Parse content
        self.parse_markdown_professional(content, doc)
        
        # Generate output filename
        output_filename = self.output_folder / f"{markdown_file.stem}_Professional.docx"
        
        # Save document
        doc.save(output_filename)
        print(f"✓ Saved: {output_filename.name}")
        
        return output_filename
    
    def convert_text_professional(self, text_file):
        """Convert text file to professional Word document"""
        
        print(f"Converting Text: {text_file.name}")
        
        # Read content with encoding fallback
        content = self.read_file_with_encoding(text_file)
        
        # Create professional Word document
        doc = Document()
        self.setup_professional_styles(doc)
        
        # Determine document title from filename
        title = self.format_filename_as_title(text_file.stem)
        
        # Add header/footer
        self.add_document_header_footer(doc, title)
        
        # Add title page
        self.add_title_page(doc, title, "Technical Specification", "Network Support Services")
        
        # Parse content
        self.parse_text_professional(content, doc, title)
        
        # Generate output filename
        output_filename = self.output_folder / f"{text_file.stem}_Professional.docx"
        
        # Save document
        doc.save(output_filename)
        print(f"✓ Saved: {output_filename.name}")
        
        return output_filename
    
    def read_file_with_encoding(self, file_path):
        """Read file with multiple encoding attempts"""
        
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, read as binary and decode with errors='ignore'
        with open(file_path, 'rb') as f:
            return f.read().decode('utf-8', errors='ignore')
    
    def setup_professional_styles(self, doc):
        """Setup professional document styles"""
        
        # Configure document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # Create custom styles
        styles = doc.styles
        
        # Custom heading styles
        try:
            # Main heading style
            main_heading_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            main_heading_style.font.name = 'Arial'
            main_heading_style.font.size = Pt(16)
            main_heading_style.font.bold = True
            main_heading_style.font.color.rgb = RGBColor(0x2F, 0x75, 0xB5)  # Blue color
            main_heading_style.paragraph_format.space_before = Pt(24)
            main_heading_style.paragraph_format.space_after = Pt(12)
            
            # Subheading style
            sub_heading_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            sub_heading_style.font.name = 'Arial'
            sub_heading_style.font.size = Pt(14)
            sub_heading_style.font.bold = True
            sub_heading_style.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)
            sub_heading_style.paragraph_format.space_before = Pt(18)
            sub_heading_style.paragraph_format.space_after = Pt(6)
            
            # Code style
            code_style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = 'Consolas'
            code_style.font.size = Pt(9)
            code_style.paragraph_format.left_indent = Inches(0.5)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
            
        except ValueError:
            # Styles already exist
            pass
    
    def add_document_header_footer(self, doc, title):
        """Add professional header and footer"""
        
        # Add header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = f"BITMARCK Network Services - {title}"
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated: {datetime.now().strftime('%B %d, %Y')} | Page "
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    def add_title_page(self, doc, title, subtitle, organization):
        """Add professional title page"""
        
        # Organization name
        org_para = doc.add_paragraph()
        org_run = org_para.add_run(organization)
        org_run.font.name = 'Arial'
        org_run.font.size = Pt(24)
        org_run.font.bold = True
        org_run.font.color.rgb = RGBColor(0x2F, 0x75, 0xB5)
        org_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add spacing
        for _ in range(3):
            doc.add_paragraph()
        
        # Main title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_run = subtitle_para.add_run(subtitle)
        subtitle_run.font.name = 'Arial'
        subtitle_run.font.size = Pt(16)
        subtitle_run.italic = True
        subtitle_run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)
        subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add spacing
        for _ in range(8):
            doc.add_paragraph()
        
        # Document info
        info_para = doc.add_paragraph()
        info_text = f"""
Document Type: Technical Specification
Project: Network Support Services Tender
Date: {datetime.now().strftime('%B %d, %Y')}
Version: 1.0
        """
        info_run = info_para.add_run(info_text.strip())
        info_run.font.name = 'Arial'
        info_run.font.size = Pt(12)
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Page break
        doc.add_page_break()
    
    def format_filename_as_title(self, filename):
        """Format filename as professional title"""
        
        # Replace common separators and format
        title = filename.replace('_', ' ').replace('-', ' ')
        title = re.sub(r'\s+', ' ', title)  # Multiple spaces to single
        
        # Capitalize properly
        words = title.split()
        formatted_words = []
        
        for word in words:
            if word.lower() in ['lot', 'network', 'support', 'services', 'cisco', 'palo', 'alto', 'f5']:
                formatted_words.append(word.upper() if word.lower() in ['lot', 'f5'] else word.title())
            else:
                formatted_words.append(word.title())
        
        return ' '.join(formatted_words)
    
    def parse_markdown_professional(self, content, doc):
        """Parse markdown with professional formatting"""
        
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('#').strip()
                
                if level == 1:
                    para = doc.add_paragraph(header_text, style='CustomHeading1')
                elif level == 2:
                    para = doc.add_paragraph(header_text, style='CustomHeading2')
                else:
                    para = doc.add_paragraph()
                    run = para.add_run(header_text)
                    run.font.bold = True
                    run.font.size = Pt(12)
                continue
            
            # Code blocks
            if line.startswith('```'):
                continue
            
            # Lists
            if line.startswith(('- ', '* ', '+ ')) or re.match(r'^\d+\.\s', line):
                list_text = re.sub(r'^[-*+]\s*|^\d+\.\s*', '', line)
                doc.add_paragraph(list_text, style='List Bullet' if line.startswith(('- ', '* ', '+ ')) else 'List Number')
                continue
            
            # Regular paragraph
            doc.add_paragraph(line)
    
    def parse_text_professional(self, content, doc, title):
        """Parse text content with professional formatting"""
        
        lines = content.split('\n')
        
        # Add table of contents placeholder
        toc_para = doc.add_paragraph("Table of Contents")
        toc_run = toc_para.runs[0]
        toc_run.font.bold = True
        toc_run.font.size = Pt(14)
        doc.add_paragraph("(Generated automatically based on document structure)")
        doc.add_paragraph()
        
        current_section = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                continue
            
            # Detect section headers (various patterns)
            if self.is_section_header(line):
                para = doc.add_paragraph(line, style='CustomHeading1')
                current_section = line
                continue
            
            # Detect subsections
            if self.is_subsection_header(line):
                para = doc.add_paragraph(line, style='CustomHeading2')
                continue
            
            # Detect lists
            if line.startswith(('- ', '* ', '• ', '○ ')):
                list_text = re.sub(r'^[-*•○]\s*', '', line)
                doc.add_paragraph(list_text, style='List Bullet')
                continue
            
            if re.match(r'^\d+[\.)]\s', line):
                list_text = re.sub(r'^\d+[\.)]\s*', '', line)
                doc.add_paragraph(list_text, style='List Number')
                continue
            
            # Questions or important items
            if line.startswith(('Question', 'Q:', 'Frage')):
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x2F, 0x75, 0xB5)
                continue
            
            # Regular paragraph
            doc.add_paragraph(line)
    
    def is_section_header(self, line):
        """Detect if line is a section header"""
        
        indicators = [
            line.isupper() and len(line) < 100,
            line.startswith(('LOT ', 'Lot ', 'SECTION', 'Section')),
            line.endswith(':') and len(line) < 80,
            re.match(r'^\d+\.\s*[A-Z]', line)
        ]
        
        return any(indicators)
    
    def is_subsection_header(self, line):
        """Detect if line is a subsection header"""
        
        indicators = [
            line.startswith(('Question', 'Frage', 'Answer', 'Antwort')),
            re.match(r'^\d+\.\d+\s', line),
            line.endswith(':') and not line.isupper() and len(line) < 60
        ]
        
        return any(indicators)

def main():
    """Main function for enhanced conversion"""
    
    source_folder = "/Users/adiscato/Python/TENDER"
    
    print("🔄 Enhanced Professional Document Converter")
    print("=" * 50)
    
    converter = EnhancedDocumentConverter(source_folder)
    
    try:
        converted_files = converter.convert_all_documents_professional()
        
        print("\n✅ Conversion Complete!")
        print(f"Successfully converted {len(converted_files)} files:")
        print()
        
        for file_path in converted_files:
            print(f"  📄 {file_path.name}")
        
        print(f"\n📁 All professional Word documents saved to:")
        print(f"   {converter.output_folder}")
        
        # File size information
        print(f"\n📊 File Information:")
        for file_path in converted_files:
            size_mb = file_path.stat().st_size / (1024 * 1024)
            print(f"  {file_path.name}: {size_mb:.2f} MB")
        
    except Exception as e:
        print(f"❌ Error during conversion: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()